from docx import Document
from reportlab.pdfgen import canvas

def generate_word(data):
    doc = Document()
    doc.add_heading('Procès-Verbal de Mise en Concordance', level=1)

    doc.add_paragraph(f"Conservation foncière : {data['conservation'][0]}")
    doc.add_paragraph(f"Titre Foncier N° : {data['titre_foncier'][0]} | Indice : {data['indice'][0]}")
    doc.add_paragraph(f"Propriété dite : {data['propriete_dite'][0]}")

    file_path = "document.docx"
    doc.save(file_path)
    return file_path
